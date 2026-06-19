from io import BytesIO

import fitz
from docx import Document
from fastapi import HTTPException, status

from app.utils.text_cleaner import clean_resume_text


def is_text_usable(text: str | None, min_length: int) -> bool:
    """Check if extracted text meets minimum usability requirements."""
    return text is not None and len(text.strip()) >= min_length


def extract_text_from_pdf(content: bytes) -> str:
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            pages = [page.get_text("text", sort=True) for page in document]
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not read PDF resume. The file may be encrypted, corrupt, or image-only.",
        ) from exc

    return "\n\n".join(page for page in pages if page.strip())


def extract_text_from_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not read DOCX resume. The file may be corrupt or unsupported.",
        ) from exc

    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="replace")


def extract_resume_text(content: bytes, extension: str) -> str:
    if extension == ".pdf":
        raw_text = extract_text_from_pdf(content)
    elif extension == ".docx":
        raw_text = extract_text_from_docx(content)
    elif extension == ".txt":
        raw_text = extract_text_from_txt(content)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Unsupported resume file type.",
        )

    cleaned_text = clean_resume_text(raw_text)
    return cleaned_text

